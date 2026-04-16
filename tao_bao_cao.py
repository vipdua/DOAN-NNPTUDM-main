from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_heading_with_color(doc, text, level, rgb_color=(0, 51, 102)):
    """Thêm heading với màu sắc"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*rgb_color)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    return heading

def add_paragraph_with_format(doc, text, bold=False, italic=False, size=11):
    """Thêm paragraph với định dạng"""
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
    return p

def add_table_with_borders(doc, data, header_color=(0, 102, 204)):
    """Thêm bảng với viền"""
    table = doc.add_table(rows=len(data) + 1, cols=len(data[0]))
    table.style = 'Light Grid Accent 1'
    
    # Thêm header
    header_cells = table.rows[0].cells
    for i, header in enumerate(data[0]):
        header_cells[i].text = str(header)
        for run in header_cells[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        # Set header background color
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), '%02x%02x%02x' % header_color)
        header_cells[i]._element.get_or_add_tcPr().append(shading_elm)
    
    # Thêm data
    for i, row_data in enumerate(data[1:], 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = str(cell_data)
    
    return table

# Tạo document
doc = Document()

# ==================== TIÊU ĐỀ ====================
title = doc.add_heading('BÁO CÁO PHÂN TÍCH HỆ THỐNG', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0, 51, 102)
    run.font.size = Pt(16)

# Subtitle
subtitle = doc.add_paragraph('Hệ Thống Quản Lý Bán Hàng Online')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(14)
    run.font.bold = True

# Thông tin tác giả
info_table = doc.add_table(rows=3, cols=2)
info_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
info_rows = [
    ('Tác giả:', 'Phạm Quốc Trung'),
    ('Mã sinh viên:', '2280616972'),
    ('Ngày hoàn thành:', 'Tháng 4 năm 2026')
]
for i, (label, value) in enumerate(info_rows):
    info_table.rows[i].cells[0].text = label
    info_table.rows[i].cells[1].text = value
    for j in range(2):
        for run in info_table.rows[i].cells[j].paragraphs[0].runs:
            run.font.size = Pt(11)

doc.add_paragraph()

# ==================== I. MỤC LỤC ====================
doc.add_heading('I. MỤC LỤC', level=1)
chapters = [
    '1. Mô tả tổng quan hệ thống',
    '2. Kiến trúc hệ thống',
    '3. Các chức năng chính',
    '4. Mô tả API Backend',
    '5. Luồng hoạt động chính',
    '6. Các kỹ thuật kiểm thử',
    '7. Kết luận'
]
for chapter in chapters:
    doc.add_paragraph(chapter, style='List Bullet')

doc.add_page_break()

# ==================== II. MỤC ĐÍCH ====================
doc.add_heading('II. MỤC ĐÍCH DỰ ÁN', level=1)
doc.add_paragraph(
    'Dự án này nhằm xây dựng một hệ thống thương mại điện tử (e-commerce) hoàn chỉnh, '
    'cho phép người dùng tìm kiếm, xem chi tiết, và mua sắm các sản phẩm trực tuyến. '
    'Hệ thống bao gồm hai phần chính: backend xử lý logic kinh doanh và frontend cung cấp '
    'giao diện người dùng thân thiện.'
)

doc.add_page_break()

# ==================== 1. MÔ TẢ TỔNG QUAN ====================
doc.add_heading('1. MÔ TẢ TỔNG QUAN HỆ THỐNG', level=1)

doc.add_heading('1.1. Giới thiệu chung', level=2)
doc.add_paragraph(
    'Hệ thống quản lý bán hàng online (NNPTUD - Ứng Dụng Mua Bán Trực Tuyến) được xây dựng '
    'trên nền tảng Node.js + Express cho backend và React cho frontend. Hệ thống hỗ trợ các '
    'tính năng cơ bản của một sàn thương mại điện tử hiện đại.'
)

doc.add_heading('1.2. Công nghệ sử dụng', level=2)
tech_data = [
    ['Thành phần', 'Công nghệ', 'Phiên bản'],
    ['Backend', 'Node.js + Express', '4.16.1'],
    ['Frontend', 'React + Vite', '19.2.4'],
    ['Routing Frontend', 'React Router', '7.14.0'],
    ['Cơ sở dữ liệu', 'MongoDB', '9.2.4'],
    ['ORM/ODM', 'Mongoose', '9.2.4'],
    ['Xác thực', 'JWT (JSON Web Token)', '9.0.3'],
    ['Mã hóa', 'Bcrypt', '6.0.0'],
    ['HTTP Client', 'Axios', '1.14.0'],
]
add_table_with_borders(doc, tech_data)

doc.add_heading('1.3. Các tính năng chính', level=2)
features = [
    'Xác thực người dùng (Register, Login, Logout)',
    'Quản lý sản phẩm (Liệt kê, Chi tiết, Tìm kiếm, Lọc giá)',
    'Quản lý danh mục sản phẩm',
    'Giỏ hàng (Thêm, Xóa, Tăng/Giảm số lượng)',
    'Quản lý người dùng (Dành cho Admin)',
    'Hệ thống phân quyền (User, Moderator, Admin)',
    'Khóa tài khoản sau nhiều lần đăng nhập sai',
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_page_break()

# ==================== 2. KIẾN TRÚC HỆ THỐNG ====================
doc.add_heading('2. KIẾN TRÚC HỆ THỐNG', level=1)

doc.add_heading('2.1. Tổng quan kiến trúc', level=2)
doc.add_paragraph(
    'Hệ thống được xây dựng theo mô hình kiến trúc ba tầng (3-Tier Architecture):'
)

architecture_items = [
    ('Presentation Layer (Frontend):', 
     'React + React Router - Cung cấp giao diện người dùng'),
    ('Business Logic Layer (Backend):', 
     'Express.js - Xử lý logic ứng dụng, xác thực, phân quyền'),
    ('Data Layer (Database):', 
     'MongoDB - Lưu trữ dữ liệu'),
]
for title, desc in architecture_items:
    p = doc.add_paragraph()
    p_run = p.add_run(title)
    p_run.bold = True
    p.add_run(desc)

doc.add_heading('2.2. Cấu trúc thư mục Backend', level=2)
backend_structure = [
    ['Thư mục/File', 'Mô tả'],
    ['app.js', 'File mặc định của Express, cấu hình middleware và routes'],
    ['bin/www', 'Điểm khởi động server'],
    ['package.json', 'Quản lý dependencies'],
    ['routes/', 'Chứa các file định nghĩa API endpoints (users, products, carts, auth...)'],
    ['schemas/', 'Chứa MongoDB schemas (User, Product, Cart, Category...)'],
    ['controllers/', 'Chứa logic xử lý dữ liệu'],
    ['utils/', 'Chứa các hàm tiện ích (authHandler, validateHandler, mailHandler...)'],
]
add_table_with_borders(doc, backend_structure, (0, 102, 204))

doc.add_heading('2.3. Cấu trúc thư mục Frontend', level=2)
frontend_structure = [
    ['Thư mục/File', 'Mô tả'],
    ['src/pages/', 'Chứa các page component (Home, Login, Products, Cart, Admin...)'],
    ['src/components/', 'Chứa reusable components (AdminRoute...)'],
    ['src/layouts/', 'Chứa layout components (MainLayout, AdminLayout)'],
    ['src/api/', 'Chứa Axios configuration để call API'],
    ['cypress/', 'Chứa test cases của Cypress'],
    ['selenium/', 'Chứa test cases của Selenium WebDriver'],
    ['testcafe/', 'Chứa test cases của TestCafe'],
]
add_table_with_borders(doc, frontend_structure, (204, 102, 0))

doc.add_heading('2.4. Luồng dữ liệu', level=2)
doc.add_paragraph('Frontend → Axios → Backend API → Express Routes → MongoDB → Response → Frontend')

doc.add_page_break()

# ==================== 3. CÁC CHỨC NĂNG CHÍNH ====================
doc.add_heading('3. CÁC CHỨC NĂNG CHÍNH', level=1)

doc.add_heading('3.1. Chức năng xác thực (Authentication)', level=2)
doc.add_paragraph(
    'Hệ thống xác thực được thực hiện thông qua JWT (JSON Web Token). '
    'Khi người dùng đăng nhập thành công, hệ thống sẽ trả về token được lưu trong localStorage.'
)
auth_data = [
    ['Chức năng', 'Mô tả', 'Quyền hạn'],
    ['Register', 'Đăng ký tài khoản mới', 'Không yêu cầu đăng nhập'],
    ['Login', 'Đăng nhập bằng username/email và password', 'Không yêu cầu đăng nhập'],
    ['Logout', 'Đăng xuất (xóa token)', 'Yêu cầu đăng nhập'],
    ['Change Password', 'Thay đổi mật khẩu', 'Yêu cầu đăng nhập'],
    ['Forgot Password', 'Gửi link reset password qua email', 'Không yêu cầu đăng nhập'],
]
add_table_with_borders(doc, auth_data)

doc.add_heading('3.2. Quản lý sản phẩm (Product Management)', level=2)
doc.add_paragraph('Cho phép xem danh sách sản phẩm, tìm kiếm, lọc theo giá, và xem chi tiết sản phẩm.')
product_data = [
    ['Chức năng', 'Mô tả', 'Xem tồn kho'],
    ['GET /products', 'Liệt kê tất cả sản phẩm, hỗ trợ filter', 'Có'],
    ['GET /products/:id', 'Xem chi tiết sản phẩm', 'Có'],
    ['POST /products', 'Tạo sản phẩm mới (Admin)', 'Có'],
    ['PUT /products/:id', 'Cập nhật sản phẩm (Admin)', 'Có'],
    ['DELETE /products/:id', 'Xóa sản phẩm (Admin)', 'Có'],
]
add_table_with_borders(doc, product_data)

doc.add_heading('3.3. Quản lý danh mục (Category Management)', level=2)
doc.add_paragraph(
    'Các sản phẩm được phân loại theo danh mục. Admin có thể tạo, chỉnh sửa, xóa danh mục.'
)

doc.add_heading('3.4. Giỏ hàng (Shopping Cart)', level=2)
doc.add_paragraph(
    'Người dùng có thể thêm, xóa, tăng/giảm số lượng sản phẩm trong giỏ hàng. '
    'Hệ thống kiểm tra tồn kho trước khi cho phép cập nhật.'
)
cart_data = [
    ['Chức năng', 'Mô tả', 'Kiểm tra tồn kho'],
    ['GET /carts', 'Lấy giỏ hàng của người dùng', 'N/A'],
    ['POST /carts/add', 'Thêm sản phẩm vào giỏ', 'Có'],
    ['POST /carts/remove', 'Xóa sản phẩm khỏi giỏ', 'Có'],
    ['POST /carts/decrease', 'Giảm số lượng sản phẩm', 'Có'],
]
add_table_with_borders(doc, cart_data)

doc.add_page_break()

# ==================== 4. MÔ TẢ API BACKEND ====================
doc.add_heading('4. MÔ TẢ API BACKEND', level=1)

doc.add_heading('4.1. API Xác thực (/api/v1/auth)', level=2)
auth_api = [
    ['Method', 'Endpoint', 'Mô tả', 'Dữ liệu đầu vào'],
    ['POST', '/api/v1/auth/register', 'Đăng ký người dùng', 'username, password, email'],
    ['POST', '/api/v1/auth/login', 'Đăng nhập', 'username, password'],
    ['GET', '/api/v1/auth/me', 'Lấy thông tin user hiện tại', 'Token (Header)'],
    ['POST', '/api/v1/auth/logout', 'Đăng xuất', 'Token (Header)'],
    ['POST', '/api/v1/auth/changepassword', 'Thay đổi mật khẩu', 'oldpassword, newpassword'],
    ['POST', '/api/v1/auth/forgotpassword', 'Gửi link reset password', 'email'],
]
add_table_with_borders(doc, auth_api)

doc.add_heading('4.2. API Người dùng (/api/v1/users)', level=2)
user_api = [
    ['Method', 'Endpoint', 'Mô tả', 'Quyền hạn'],
    ['GET', '/api/v1/users', 'Liệt kê tất cả người dùng', 'ADMIN'],
    ['GET', '/api/v1/users/:id', 'Xem chi tiết người dùng', 'ADMIN, MODERATOR'],
    ['POST', '/api/v1/users', 'Tạo người dùng mới', 'ADMIN'],
    ['PUT', '/api/v1/users/:id', 'Cập nhật người dùng', 'ADMIN'],
    ['DELETE', '/api/v1/users/:id', 'Xóa người dùng (soft delete)', 'ADMIN'],
]
add_table_with_borders(doc, user_api)

doc.add_heading('4.3. API Sản phẩm (/api/v1/products)', level=2)
products_api = [
    ['Method', 'Endpoint', 'Mô tả', 'Quyền hạn'],
    ['GET', '/api/v1/products', 'Liệt kê sản phẩm (hỗ trợ filter)', 'Công khai'],
    ['GET', '/api/v1/products/:id', 'Xem chi tiết sản phẩm', 'Công khai'],
    ['POST', '/api/v1/products', 'Tạo sản phẩm mới', 'ADMIN'],
    ['PUT', '/api/v1/products/:id', 'Cập nhật sản phẩm', 'ADMIN'],
    ['DELETE', '/api/v1/products/:id', 'Xóa sản phẩm (soft delete)', 'ADMIN'],
]
add_table_with_borders(doc, products_api)

doc.add_heading('4.4. API Giỏ hàng (/api/v1/carts)', level=2)
carts_api = [
    ['Method', 'Endpoint', 'Mô tả', 'Quyền hạn'],
    ['GET', '/api/v1/carts', 'Lấy giỏ hàng của user', 'Đã đăng nhập'],
    ['POST', '/api/v1/carts/add', 'Thêm sản phẩm vào giỏ', 'Đã đăng nhập'],
    ['POST', '/api/v1/carts/remove', 'Xóa sản phẩm khỏi giỏ', 'Đã đăng nhập'],
    ['POST', '/api/v1/carts/decrease', 'Giảm số lượng sản phẩm', 'Đã đăng nhập'],
]
add_table_with_borders(doc, carts_api)

doc.add_heading('4.5. Phân quyền (Authorization)', level=2)
doc.add_paragraph(
    'Hệ thống sử dụng quyền hạn dựa trên Roles:'
)
roles_data = [
    ['Role', 'Mô tả', 'Quyền hạn'],
    ['USER', 'Người dùng thông thường', 'Xem sản phẩm, mua hàng'],
    ['MODERATOR', 'Người kiểm duyệt', 'Xem người dùng, moderator forums'],
    ['ADMIN', 'Quản trị viên', 'Toàn quyền'],
]
add_table_with_borders(doc, roles_data)

doc.add_page_break()

# ==================== 5. LUỒNG HOẠT ĐỘNG CHÍNH ====================
doc.add_heading('5. LUỒNG HOẠT ĐỘNG CHÍNH', level=1)

doc.add_heading('5.1. Luồng đăng nhập (Login Flow)', level=2)
login_steps = [
    ('Bước 1:', 'Người dùng nhập email/username và mật khẩu'),
    ('Bước 2:', 'Frontend gửi request POST /api/v1/auth/login'),
    ('Bước 3:', 'Backend kiểm tra username có tồn tại không'),
    ('Bước 4:', 'Backend kiểm tra mật khẩu bằng bcrypt'),
    ('Bước 5:', 'Nếu đúng, tạo JWT token với thời hạn 24 giờ'),
    ('Bước 6:', 'Lưu token vào cookie và localStorage'),
    ('Bước 7:', 'Điều hướng dựa trên role (Admin → /admin/dashboard, User → /)'),
]
for step, desc in login_steps:
    p = doc.add_paragraph()
    p_run = p.add_run(step)
    p_run.bold = True
    p.add_run(desc)

doc.add_heading('5.2. Luồng mua hàng (Shopping Flow)', level=2)
shopping_steps = [
    ('Bước 1:', 'Người dùng xem danh sách sản phẩm'),
    ('Bước 2:', 'Người dùng có thể lọc theo giá hoặc danh mục'),
    ('Bước 3:', 'Người dùng xem chi tiết sản phẩm và tồn kho'),
    ('Bước 4:', 'Người dùng thêm sản phẩm vào giỏ hàng'),
    ('Bước 5:', 'Backend kiểm tra số lượng tồn kho'),
    ('Bước 6:', 'Nếu có đủ tồn kho, thêm vào giỏ hàng'),
    ('Bước 7:', 'Người dùng có thể tăng/giảm số lượng hoặc xóa sản phẩm'),
    ('Bước 8:', 'Người dùng tiến hành thanh toán (chức năng sẽ được phát triển)'),
]
for step, desc in shopping_steps:
    p = doc.add_paragraph()
    p_run = p.add_run(step)
    p_run.bold = True
    p.add_run(desc)

doc.add_heading('5.3. Sơ đồ quy trình (Flow Diagram)', level=2)
doc.add_paragraph('Sơ đồ quy trình từ Login đến Checkout:')
flowchart = '''
START
  ↓
[Người dùng truy cập trang]
  ↓
{Đã đăng nhập?} → NO → [Trang Login]
  ↓ YES             ↓
  ↓          [Login/Register]
  ↓                 ↓
[Trang chủ] ← ← ← ← ↑
  ↓
[Xem sản phẩm → Lọc/Tìm kiếm]
  ↓
[Click vào sản phẩm]
  ↓
[Xem chi tiết + Tồn kho]
  ↓
{Thêm vào giỏ?} → NO → [Trang sản phẩm]
  ↓ YES               ↑
  ↓            {Tiếp tục?}
[Cập nhật giỏ hàng]  ↑
  ↓                 YES
  ↓                  ↓
[Xem Giỏ hàng] << ← [Xem sản phẩm]
  ↓
{Thay đổi số lượng?}
  ├─ YES → [Tăng/Giảm]
  └─ NO  → (tiếp tục)
  ↓
[Thanh toán] (Sẽ phát triển)
  ↓
END
'''
doc.add_paragraph(flowchart, style='List')

doc.add_page_break()

# ==================== 6. CÁC KỸ THUẬT KIỂM THỬ ====================
doc.add_heading('6. CÁC KỸ THUẬT KIỂM THỬ', level=1)

doc.add_heading('6.1. Tổng quan về kiểm thử', level=2)
doc.add_paragraph(
    'Dự án sử dụng ba công cụ kiểm thử tự động để đảm bảo chất lượng: '
    'Cypress, TestCafe, và Selenium WebDriver.'
)

doc.add_heading('6.2. Cypress Testing', level=2)
doc.add_paragraph(
    'Cypress là framework kiểm thử end-to-end được xây dựng trên Node.js, '
    'cho phép kiểm thử trực tiếp trên trình duyệt.'
)
cypress_data = [
    ['Tính năng', 'Mô tả'],
    ['Location', 'frontend/cypress/e2e/'],
    ['Test File', 'auth.cy.js'],
    ['Framework', 'Mocha (BDD)'],
    ['Selectors', 'CSS selectors (cy.get)'],
    ['Test Cases', 'Login success, login failed, boundary testing'],
]
add_table_with_borders(doc, cypress_data)

doc.add_heading('Ví dụ Test Case Cypress:', level=3)
doc.add_paragraph('Login thành công với username hợp lệ:')
cypress_example = '''
it('TC01 - Login thành công', () => {
    cy.get('input').eq(0).type('admin');
    cy.get('input').eq(1).type('Admin123@');
    cy.get('button[type="submit"]').click();
    cy.url().should('include', '/admin/dashboard');
});
'''
doc.add_paragraph(cypress_example, style='List')

doc.add_heading('6.3. TestCafe Framework', level=2)
doc.add_paragraph(
    'TestCafe là framework kiểm thử code-free và cloud-ready, '
    'hỗ trợ kiểm thử trên nhiều trình duyệt.'
)
testcafe_data = [
    ['Tính năng', 'Mô tả'],
    ['Location', 'frontend/testcafe/'],
    ['Test File', 'category.test.js'],
    ['Framework', 'TestCafe'],
    ['Browser Support', 'Chrome, Firefox, Safari, Edge'],
    ['Chạy Test', 'testcafe chrome category.test.js'],
]
add_table_with_borders(doc, testcafe_data)

doc.add_heading('6.4. Selenium WebDriver', level=2)
doc.add_paragraph(
    'Selenium WebDriver là công cụ kiểm thử độc lập, '
    'hỗ trợ kiểm thử trên nhiều nền tảng.'
)
selenium_data = [
    ['Tính năng', 'Mô tả'],
    ['Location', 'frontend/selenium/'],
    ['Test File', 'cart.test.cjs'],
    ['Driver', 'ChromeDriver'],
    ['Language', 'JavaScript (Node.js)'],
    ['Package', 'selenium-webdriver v4.43.0'],
]
add_table_with_borders(doc, selenium_data)

doc.add_heading('6.5. Kỹ thuật kiểm thử được sử dụng', level=2)
testing_techniques = [
    ('Equivalence Partitioning:', 
     'Chia input thành các nhóm có cùng hành vi - VD: Test login với username hợp lệ/không hợp lệ'),
    ('Boundary Value Testing:', 
     'Kiểm thử tại biên của input - VD: Password cực ngắn, username trống'),
    ('Positive Testing:', 
     'Kiểm thử với dữ liệu hợp lệ - VD: Login với credentials đúng'),
    ('Negative Testing:', 
     'Kiểm thử với dữ liệu không hợp lệ - VD: Login với password sai'),
]
for technique, description in testing_techniques:
    p = doc.add_paragraph()
    p_run = p.add_run(technique)
    p_run.bold = True
    p.add_run('\n' + description)

doc.add_heading('6.6. Test Cases mẫu (từ auth.cy.js)', level=2)
test_cases = [
    ('TC01', 'Login thành công', 'Username: admin, Password: Admin123@'),
    ('TC02', 'Sai username', 'Username: saiuser, Password: Admin123@'),
    ('TC03', 'Sai password', 'Username: admin, Password: sai123'),
    ('TC04', 'Bỏ trống form', 'Gửi form không điền dữ liệu'),
    ('TC05', 'Password cực ngắn', 'Username: admin, Password: 1'),
]
for tc_id, description, data in test_cases:
    p = doc.add_paragraph()
    p_run = p.add_run(f'{tc_id}: ')
    p_run.bold = True
    p.add_run(f'{description} - {data}')

doc.add_page_break()

# ==================== 7. MÔ TẢ CÁC SCHEMA DATABASEE ====================
doc.add_heading('7. MÔ TẢ CÁC SCHEMA DATABASE', level=1)

doc.add_heading('7.1. User Schema', level=2)
schemas_explanation = '''
Users: Lưu trữ thông tin người dùng
- _id: MongoDB ObjectId (tự động tạo)
- username: String (unique, required)
- password: String (mã hóa bcrypt, required)
- email: String (unique, required)
- fullName: String (tùy chọn)
- avatarUrl: String (mặc định: avatar placeholder)
- status: Boolean (active/inactive)
- role: ObjectId (tham chiếu Roles collection)
- loginCount: Number (tính số lần đăng nhập sai)
- lockTime: Date (thời gian khóa tài khoản)
- forgotPasswordToken: String (token reset password)
- forgotPasswordTokenExp: Date (hạn sử dụng token)
- isDeleted: Boolean (soft delete)
- timestamps: createdAt, updatedAt (tự động)
'''
doc.add_paragraph(schemas_explanation)

doc.add_heading('7.2. Product Schema', level=2)
product_schema = '''
Products: Lưu trữ thông tin sản phẩm
- _id: MongoDB ObjectId
- title: String (unique, required)
- slug: String (unique - để SEO friendly URLs)
- price: Number (giá sản phẩm)
- description: String (mô tả chi tiết)
- images: Array of String (danh sách ảnh)
- category: ObjectId (tham chiếu Categories)
- isDeleted: Boolean (soft delete)
- timestamps: createdAt, updatedAt
'''
doc.add_paragraph(product_schema)

doc.add_heading('7.3. Cart Schema', level=2)
cart_schema = '''
Carts: Lưu trữ giỏ hàng của người dùng
- _id: MongoDB ObjectId
- user: ObjectId (tham chiếu Users - unique)
- products: Array của SubSchema ItemCart
  - product: ObjectId (tham chiếu Products)
  - quantity: Number (số lượng)
'''
doc.add_paragraph(cart_schema)

doc.add_heading('7.4. Inventory Schema', level=2)
inventory_schema = '''
Inventories: Lưu trữ thông tin tồn kho
- _id: MongoDB ObjectId
- product: ObjectId (tham chiếu Products - unique)
- stock: Number (số lượng tồn kho)
- timestamps: createdAt, updatedAt
'''
doc.add_paragraph(inventory_schema)

doc.add_page_break()

# ==================== 8. PHÂN TÍCH HỆ THỐNG QUẢN LÝ PHIÊN ====================
doc.add_heading('8. PHÂN TÍCH HỆ THỐNG QUẢN LÝ PHIÊN', level=1)

doc.add_heading('8.1. Xác thực người dùng (Authentication)', level=2)
doc.add_paragraph(
    'Hệ thống sử dụng JWT (JSON Web Token) cho xác thực:'
)
jwt_desc = '''
- Khi login thành công, backend tạo JWT token với thời hạn 24 giờ
- Token được lưu trong localStorage (browser) và cookie (httpOnly, secure)
- Mỗi request đến API được gửi kèm token trong header Authorization
- Backend xác thực token bằng middleware CheckLogin
- Nếu token hộp lệ, req.user được gán và tiếp tục xử lý
- Nếu token không hợp lệ hoặc hết hạn, trả về lỗi 401
'''
doc.add_paragraph(jwt_desc)

doc.add_heading('8.2. Phân quyền (Authorization)', level=2)
doc.add_paragraph(
    'Sau khi xác thực, hệ thống kiểm tra quyền hạn dựa trên role:'
)
auth_desc = '''
- Middleware CheckRole('ADMIN', 'MODERATOR') kiểm tra role của user
- Nếu user không có quyền hạn, trả về lỗi 403 Forbidden
- Admin có toàn quyền: tạo, sửa, xóa users, products, categories
- User thông thường chỉ có thể xem sản phẩm và quản lý giỏ hàng
- Moderator có thể xem thông tin users
'''
doc.add_paragraph(auth_desc)

doc.add_heading('8.3. Bảo mật mật khẩu', level=2)
doc.add_paragraph(
    'Mật khẩu được bảo vệ bằng bcrypt:'
)
password_desc = '''
- Trước khi lưu vào DB, bcrypt tạo salt (10 rounds)
- Mật khẩu được hash với salt, tạo ra hash String không thể reverse
- Khi login, bcrypt so sánh password nhập vào với hash lưu trữ
- Nếu bcrypt.compareSync() trả về true, mật khẩu đúng
- Nếu password sai, hệ thống tăng loginCount
- Nếu loginCount >= 3, tài khoản bị khóa 1 giờ (khóa tài khoản có thể bật/tắt)
'''
doc.add_paragraph(password_desc)

doc.add_page_break()

# ==================== 9. KÊTU LUẬN ====================
doc.add_heading('9. KẾT LUẬN', level=1)

doc.add_heading('9.1. Những điểm nổi bật', level=2)
highlights = [
    'Kiến trúc hệ thống rõ ràng, dễ bảo trì với 3-tier architecture',
    'Sử dụng JWT cho xác thực an toàn và hiệu quả',
    'Hệ thống phân quyền linh hoạt với 3 roles (User, Moderator, Admin)',
    'Kiểm tra tồn kho trước khi cho phép thêm vào giỏ hàng',
    'Sử dụng 3 công cụ kiểm thử: Cypress, TestCafe, Selenium',
    'Bảo vệ mật khẩu bằng bcrypt (hash 10 rounds)',
    'Soft delete để bảo toàn dữ liệu lịch sử',
    'Khóa tài khoản sau nhiều lần đăng nhập sai',
]
for highlight in highlights:
    doc.add_paragraph(highlight, style='List Bullet')

doc.add_heading('9.2. Những điểm có thể cải thiện', level=2)
improvements = [
    'Thêm chức năng thanh toán online (Stripe, PayPal)',
    'Thêm chức năng review sản phẩm từ khách hàng',
    'Thêm chức năng wishlist (yêu thích sản phẩm)',
    'Tối ưu hóa tỷ lệ xoay vòng của JWT token',
    'Thêm logging chi tiết để theo dõi hoạt động',
    'Thêm rate limiting để chống DDoS',
    'Tách file lớn thành các middleware nhỏ hơn',
    'Thêm unit tests cho backend',
    'Triển khai caching (Redis) để cải thiện hiệu suất',
    'Thêm chức năng tìm kiếm nâng cao (full-text search)',
]
for improvement in improvements:
    doc.add_paragraph(improvement, style='List Bullet')

doc.add_heading('9.3. Tổng kết', level=2)
doc.add_paragraph(
    'Dự án này đã xây dựng thành công một hệ thống thương mại điện tử '
    'với các chức năng cơ bản hoàn chỉnh. Sử dụng công nghệ hiện đại '
    '(Node.js, React, MongoDB) kết hợp với các best practices như JWT '
    'authentication, role-based authorization, và comprehensive testing. '
    'Dự án có tiềm năng phát triển thêm nhiều tính năng nâng cao để '
    'trở thành một nền tảng thương mại điện tử chuyên nghiệp.'
)

doc.add_page_break()

# ==================== PHỤ LỤC ====================
doc.add_heading('PHỤ LỤC: HƯỚNG DẪN RUN DỰ ÁN', level=1)

doc.add_heading('A. Cài đặt Backend', level=2)
backend_setup = '''
1. Di chuyển đến thư mục gốc dự án:
   cd DOAN-NNPTUDM-main

2. Cài đặt dependencies:
   npm install

3. Chạy seeder để tạo dữ liệu ban đầu:
   node seeder.js
   (Tạo admin user, roles, sample products, categories)

4. Khởi động server:
   npm start
   Server sẽ chạy tại http://localhost:3000 (hoặc port khác)
   
5. Kết nối MongoDB:
   - Mặc định kết nối local MongoDB tại mongodb://127.0.0.1:27017/SangT4
   - Có thể thay đổi trong app.js
'''
doc.add_paragraph(backend_setup)

doc.add_heading('B. Cài đặt Frontend', level=2)
frontend_setup = '''
1. Di chuyển đến thư mục frontend:
   cd frontend

2. Cài đặt dependencies:
   npm install

3. Cấu hình API URL trong src/api/axios.js:
   - Đảm bảo baseURL trỏ đến backend (default: http://localhost:3000)

4. Chạy development server:
   npm run dev
   Frontend sẽ chạy tại http://localhost:5173

5. Build cho production:
   npm run build
'''
doc.add_paragraph(frontend_setup)

doc.add_heading('C. Chạy các Test', level=2)
testing_setup = '''
1. Cypress End-to-End Tests:
   cd frontend
   npx cypress open
   - Chọn E2E Testing
   - Chọn Chrome hoặc Firefox
   - Chọn file auth.cy.js

2. TestCafe Tests:
   cd frontend
   testcafe chrome testcafe/category.test.js

3. Selenium WebDriver Tests:
   cd frontend
   node selenium/cart.test.cjs
   - Yêu cầu ChromeDriver cài đặt sẵn
'''
doc.add_paragraph(testing_setup)

doc.add_heading('D. Tài khoản Admin Mặc định', level=2)
doc.add_paragraph('Username: admin')
doc.add_paragraph('Password: Admin123@')

# Lưu document
output_path = 'BAO_CAO_PHAN_TICH_HE_THONG.docx'
doc.save(output_path)
print(f'Báo cáo đã được tạo: {output_path}')
